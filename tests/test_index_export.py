from __future__ import annotations

import contextlib
import os
import tempfile
import unittest
from pathlib import Path

from docx import Document

from src.me_finder.indexer import build_index


@contextlib.contextmanager
def _corpus_cwd():
    previous = os.getcwd()
    temp_dir = tempfile.mkdtemp()
    try:
        root = Path(temp_dir)
        corpus = root / "corpus" / "raw_docx"
        corpus.mkdir(parents=True)
        document = Document()
        document.add_paragraph("马克思恩格斯文集 第1卷")
        document.add_paragraph("宗教是人民的鸦片。")
        document.save(str(corpus / "马克思恩格斯文集_第01卷.docx"))
        os.chdir(root)
        yield root, corpus
    finally:
        os.chdir(previous)


class IndexExportTests(unittest.TestCase):
    def test_json_export_is_off_by_default_sqlite_always_built(self) -> None:
        with _corpus_cwd() as (root, corpus):
            json_path = root / "data" / "index.json"
            db_path = root / "data" / "index.sqlite3"
            build_index(corpus_dir=corpus, index_path=json_path, database_path=db_path)
            self.assertFalse(json_path.exists(), "JSON backup must not be written by default")
            self.assertTrue(db_path.exists(), "SQLite index must always be built")

    def test_json_export_written_when_opted_in(self) -> None:
        with _corpus_cwd() as (root, corpus):
            json_path = root / "data" / "index.json"
            db_path = root / "data" / "index.sqlite3"
            build_index(corpus_dir=corpus, index_path=json_path, database_path=db_path, export_json=True)
            self.assertTrue(json_path.exists(), "JSON backup must be written when export_json=True")
            self.assertTrue(db_path.exists())


if __name__ == "__main__":
    unittest.main()
